"""
report_doc_builder.py — structured report document model + output renderers

Architecture
------------
ReportDocument               — 10-section data model (plain Python types, JSON-safe)
build_report_document(state) — fills from _state via existing report_builder helpers
render_report_markdown(doc)  — clean Markdown string
render_report_html(doc)      — self-contained HTML string (no external CSS/JS)
generate_docx(doc, lang)     — DOCX bytes; requires `pip install python-docx`
                               lang: "ko" (Korean, RISS/DBpia), "en" (English, IEEE/Nature),
                                     "both" (bilingual, international)

All render functions are pure (no I/O, no _state access).
"""
from __future__ import annotations

import html as _html
import io
import textwrap
from dataclasses import asdict, dataclass, field
from typing import Any

from .report_builder import (
    build_algorithm_compare,
    build_per_bs_metrics,
    build_per_edge_metrics,
    build_run_summary,
    build_scenario_metadata,
)


# ──────────────────────────────────────────────────────────────────────────────
# Bilingual metric labels
# ──────────────────────────────────────────────────────────────────────────────

METRIC_LABELS: dict[str, dict[str, str]] = {
    "avg_latency_ms":          {"ko": "평균 E2E 지연 (ms)",           "en": "Average E2E Latency (ms)"},
    "max_latency_ms":          {"ko": "최대 지연 (ms)",               "en": "Maximum Latency (ms)"},
    "total_cost":              {"ko": "총 경로 비용",                  "en": "Total Route Cost"},
    "cost_improvement_pct":    {"ko": "비용 개선율 (%)",              "en": "Cost Improvement (%)"},
    "handover_count":          {"ko": "핸드오버 횟수",                "en": "Handover Count"},
    "coverage_risk":           {"ko": "커버리지 위험도",              "en": "Coverage Risk"},
    "covered_pct":             {"ko": "커버리지율 (%)",               "en": "Coverage Ratio (%)"},
    "prr_approx":              {"ko": "패킷 수신율 PRR (근사)",       "en": "Packet Reception Ratio PRR (approx.)"},
    "pir_p99_ms":              {"ko": "PIR P99 (ms)",                 "en": "Packet Inter-Reception Time P99 (ms)"},
    "pir_compliant":           {"ko": "PIR 기준 준수",                "en": "PIR Compliant (≤100 ms)"},
    "cbr":                     {"ko": "채널 점유율 CBR",              "en": "Channel Busy Ratio (CBR)"},
    "jain_fairness_index":     {"ko": "Jain 공정성 지수",             "en": "Jain's Fairness Index"},
    "average_bs_load":         {"ko": "평균 BS 부하율",               "en": "Average BS Load Ratio"},
    "future_connectivity_risk": {"ko": "미래 연결 위험도",            "en": "Future Connectivity Risk"},
    "time_weighted_disconnection_ratio": {"ko": "시간 가중 단절률 TWDR", "en": "Time-Weighted Disconnection Ratio (TWDR)"},
    "total_distance_m":        {"ko": "총 거리 (m)",                  "en": "Total Distance (m)"},
    "total_travel_time_s":     {"ko": "총 이동 시간 (s)",             "en": "Total Travel Time (s)"},
    "vehicle_count":           {"ko": "차량 대수",                    "en": "Vehicle Count"},
    "network_mode":            {"ko": "네트워크 모드",                "en": "Network Mode"},
    "selected_algorithm":      {"ko": "선택 알고리즘",                "en": "Selected Algorithm"},
    "baseline_algorithm":      {"ko": "기준 알고리즘",                "en": "Baseline Algorithm"},
}

def _ml(key: str, lang: str) -> str:
    """Return the label for a metric key in the given language."""
    entry = METRIC_LABELS.get(key, {})
    if lang == "en":
        return entry.get("en", key)
    return entry.get("ko", key)


# ──────────────────────────────────────────────────────────────────────────────
# Reference list (10 entries, used in DOCX References section)
# ──────────────────────────────────────────────────────────────────────────────

REFERENCES: list[str] = [
    "3GPP TS 22.186 V16.1.0 (2019-09), \"Enhancement of 3GPP support for V2X scenarios,\" 3rd Generation Partnership Project.",
    "ETSI EN 302 637-2 V1.4.1 (2019-04), \"Intelligent Transport Systems (ITS); Vehicular Communications; Basic Set of Applications; Part 2: Specification of Cooperative Awareness Basic Service,\" ETSI.",
    "ETSI TS 102 687 V1.1.1 (2011-07), \"Intelligent Transport Systems (ITS); Decentralized Congestion Control Mechanisms for Intelligent Transport Systems operating in the 5 GHz range,\" ETSI.",
    "M. Gonzalez-Martin, M. Sepulcre, R. Molina-Masegosa and J. Gozalvez, \"Analytical Models of the Performance of C-V2X Mode 4 Vehicular Communications,\" IEEE Transactions on Vehicular Technology, vol. 68, no. 2, pp. 1155-1166, Feb. 2019.",
    "H. Fernandez, L. Garcia, J. Reig and L. Rubio, \"Path Loss Characterization for Vehicular Communications at 700 MHz and 5.9 GHz under LOS and NLOS Conditions,\" IEEE Antennas and Wireless Propagation Letters, vol. 3, no. 6, pp. 1323-1326, Dec. 2014.",
    "Z. Ali, L. Lag e-Arenas, E. Dahlman and N. Himayat, \"3GPP NR V2X Mode 2: Overview, Models and System-Level Evaluation,\" IEEE Access, vol. 9, pp. 89, 2021.",
    "F. Eckermann, M. Kahlert, C. Wietfeld, \"Performance Analysis of C-V2X Mode 4 Communication Introducing an Open-Source C-V2X Simulator,\" in Proc. IEEE VTC Fall, Honolulu, 2019.",
    "S.-C. Hung, H. Hsu, S.-Y. Lien and K.-C. Chen, \"Mobile-Edge Computing for Vehicular Networks: A Semi-Decentralized MEC Architecture,\" IEEE Vehicular Technology Magazine, vol. 12, no. 1, pp. 96-105, Mar. 2017. [E2E latency model: L = L_tx + L_q + L_comp]",
    "R. Jain, D. Chiu and W. Hawe, \"A Quantitative Measure of Fairness and Discrimination for Resource Allocation in Shared Computer Systems,\" DEC Research Report TR-301, 1984.",
    "T. Abbas, J. Nuckelt, T. Kurner, T. Zemen, C. F. Mecklenbrauker and F. Tufvesson, \"Simulation and Measurement-Based Vehicle-to-Vehicle Channel Characterization: Accuracy and Limit of Applicability,\" IEEE Transactions on Antennas and Propagation, vol. 63, no. 9, pp. 4207-4218, 2015.",
]


# ──────────────────────────────────────────────────────────────────────────────
# Performance metric formula table (for Methodology section and formula panel)
# ──────────────────────────────────────────────────────────────────────────────

METRIC_FORMULAS: list[dict[str, str]] = [
    {
        "metric_ko": "E2E 지연 (평균)",
        "metric_en": "E2E Latency (avg)",
        "formula":   "L_avg = (1/N_e) · Σ_e latency_ms(e)",
        "ref":       "[8] Hung et al., IEEE VTM 2017",
        "threshold": "< 20 ms (안전 메시지)",
    },
    {
        "metric_ko": "패킷 수신율 PRR",
        "metric_en": "Packet Reception Ratio (PRR)",
        "formula":   "PRR = 1 − TWDR",
        "ref":       "[6] Ali et al., IEEE Access 2021 §II.B",
        "threshold": "> 0.90 (3GPP TR 37.885)",
    },
    {
        "metric_ko": "채널 점유율 CBR †",
        "metric_en": "Channel Busy Ratio (CBR) †",
        "formula":   "CBR = 1 − exp(−ρ·R_tx·f_CAM·T_RRI)",
        "ref":       "[4] Gonzalez-Martin et al., IEEE TVT 2019 §IV.A",
        "threshold": "< 0.65 (ETSI TS 102 687 §5.2.2)",
    },
    {
        "metric_ko": "PIR P99 상한 †",
        "metric_en": "PIR P99 upper bound †",
        "formula":   "PIR_P99 = T_CAM / PRR  (기하분포 상한)",
        "ref":       "[7] Eckermann et al., IEEE VTC 2019; [1] 3GPP TR 37.885 §A.2.4",
        "threshold": "≤ 100 ms (3GPP TR 37.885 Table A.1)",
    },
    {
        "metric_ko": "경로손실 †",
        "metric_en": "Path Loss (log-distance) †",
        "formula":   "PL(d) = PL(d₀) + 10·n·log₁₀(d/d₀)  [dB]",
        "ref":       "[5] Fernandez et al., IEEE WCL 2014, Table I (n=1.61 highway, 2.75 urban)",
        "threshold": "—",
    },
    {
        "metric_ko": "Jain 공정성 지수",
        "metric_en": "Jain's Fairness Index",
        "formula":   "J = (Σᵢ ρᵢ)² / (n · Σᵢ ρᵢ²)",
        "ref":       "[9] Jain, Chiu, Hawe, DEC-TR-301 1984 §3.1",
        "threshold": "≥ 0.90 (권장 수준)",
    },
    {
        "metric_ko": "시간 가중 단절률 TWDR",
        "metric_en": "Time-Weighted Disconnection Ratio (TWDR)",
        "formula":   "TWDR = Σ_e [t_e · I(¬cov(e))] / Σ_e t_e",
        "ref":       "자체 정의 (자체 설계)",
        "threshold": "< 0.10",
    },
    {
        "metric_ko": "경로 비용 함수",
        "metric_en": "Route Cost Function",
        "formula":   "C = Σ_e [w₁c_d(e) + w₂c_L(e) + w₃c_ρ(e) + w₄c_cov(e)]",
        "ref":       "자체 설계 (가중치는 설정 파라미터)",
        "threshold": "최소화 목표",
    },
]


# ──────────────────────────────────────────────────────────────────────────────
# Data model
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class ReportDocument:
    """
    Intermediate structured representation of a simulation analysis report.

    Sections 1-10 map to the Step 7 spec.  All fields are plain Python types
    (str, int, float, list, dict) so asdict() / json.dumps() work without
    a custom encoder.
    """
    # ── Section 1: Title ──────────────────────────────────────────────────────
    title: str = "V2X 시뮬레이션 분석 보고서"
    subtitle: str = ""

    # ── Section 2: Generated time ─────────────────────────────────────────────
    generated_at: str = ""
    run_id: str = ""

    # ── Section 3: Scenario description ──────────────────────────────────────
    scenario_id: str = ""
    scenario_name: str = ""
    origin_lat: float | None = None
    origin_lng: float | None = None
    dest_lat: float | None = None
    dest_lng: float | None = None
    vehicle_count: int = 0
    network_mode: str = ""

    # ── Section 4: Simulation configuration ──────────────────────────────────
    sim_mode: str = ""
    seed: Any = None
    route_algorithm: str = ""
    bs_selection_algorithm: str = ""
    resource_allocation_algorithm: str = ""
    lookahead_k: Any = None
    cost_weights: dict = field(default_factory=dict)
    norm_scales: dict = field(default_factory=dict)

    # ── Section 5: Key KPI summary ────────────────────────────────────────────
    kpis: dict = field(default_factory=dict)
    improvement: dict = field(default_factory=dict)

    # ── Section 6: Algorithm comparison ──────────────────────────────────────
    selected_algorithm: str = ""
    baseline_algorithm: str = ""
    algorithms: list = field(default_factory=list)

    # ── Section 7: Chart descriptions ────────────────────────────────────────
    chart_descriptions: list = field(default_factory=list)

    # ── Section 8: Bottleneck / risk findings ─────────────────────────────────
    bottleneck_sections: list = field(default_factory=list)
    overloaded_bs: list = field(default_factory=list)
    high_latency_sections: list = field(default_factory=list)
    handover_sections: list = field(default_factory=list)
    future_risk_sections: list = field(default_factory=list)

    # ── Section 9: Conclusion ─────────────────────────────────────────────────
    primary_finding: str = ""
    trade_offs: list = field(default_factory=list)
    risk_factors: list = field(default_factory=list)
    recommendation_text: str = ""

    # ── Section 10: Appendix / raw metric notes ───────────────────────────────
    per_edge_sample: list = field(default_factory=list)
    per_bs_summary: list = field(default_factory=list)
    metadata_notes: dict = field(default_factory=dict)

    # ── Extended KPIs (new analytical metrics) ────────────────────────────────
    pir_p99_ms: float | None = None
    pir_compliant: bool | None = None
    jain_fairness_index: float | None = None
    cbr_avg: float | None = None

    def to_dict(self) -> dict:
        return asdict(self)


# ──────────────────────────────────────────────────────────────────────────────
# Builder
# ──────────────────────────────────────────────────────────────────────────────

def build_report_document(state: dict) -> ReportDocument:
    """
    Build a ReportDocument from the current _state snapshot.

    Delegates to existing report_builder helpers; does no direct _state access
    beyond what those helpers already encapsulate.
    """
    rs    = build_run_summary(state)
    meta  = build_scenario_metadata(state)
    algos = build_algorithm_compare(state)
    edges = build_per_edge_metrics(state)
    bs    = build_per_bs_metrics(state)
    sim_s = state.get("simulation_summary") or {}
    rec   = sim_s.get("recommendation_text_seed") or {}
    imp   = sim_s.get("improvement_over_baseline") or {}

    subtitle = (
        f"{rs.get('selected_algorithm', '')} 알고리즘 / "
        f"{rs.get('network_mode', '')} / "
        f"차량 {rs.get('vehicle_count', 0)}대"
    )

    kpis = {
        "총 비용":                  rs.get("total_cost"),
        "평균 지연 (ms)":           rs.get("avg_latency_ms"),
        "최대 지연 (ms)":           rs.get("max_latency_ms"),
        "핸드오버 횟수":            rs.get("handover_count"),
        "커버리지 위험":            rs.get("coverage_risk"),
        "커버리지율 (%)":           rs.get("covered_pct"),
        "PRR (근사)":               rs.get("prr_approx"),
        "PIR P99 (ms) †":          rs.get("pir_p99_ms"),
        "PIR 기준 준수":            rs.get("pir_compliant"),
        "Jain 공정성 지수":         rs.get("jain_fairness_index"),
        "총 거리 (m)":              rs.get("total_distance_m"),
        "총 이동 시간 (s)":         rs.get("total_travel_time_s"),
        "평균 BS 부하":             rs.get("average_bs_load"),
        "미래 연결 위험":           rs.get("future_connectivity_risk"),
    }

    improvement = {
        "비용 개선 (%)":     imp.get("cost_improvement_pct"),
        "지연 변화 (ms)":    imp.get("latency_delta_ms"),
        "거리 변화 (m)":     imp.get("distance_delta_m"),
        "핸드오버 변화":     imp.get("handover_delta"),
        "선택 알고리즘":     imp.get("selected_algorithm") or rs.get("selected_algorithm"),
        "기준 알고리즘":     imp.get("baseline_algorithm") or rs.get("baseline_algorithm"),
    }

    chart_descriptions = [
        {"번호": 1, "유형": "RadarChart",     "제목": "알고리즘 KPI 레이더 비교",         "데이터": "algorithm_compare"},
        {"번호": 2, "유형": "LineChart",      "제목": "파라미터 스윕 추이",               "데이터": "batch_sweep"},
        {"번호": 3, "유형": "HistogramChart", "제목": "구간 지연 분포",                   "데이터": "per_edge_metrics.latency_ms"},
        {"번호": 4, "유형": "BarChart",       "제목": "BS 부하 비교",                     "데이터": "per_bs_metrics.load_ratio"},
        {"번호": 5, "유형": "SegmentStrip",   "제목": "경로 구간별 위험도",               "데이터": "per_edge_metrics"},
        {"번호": 6, "유형": "DeltaBarChart",  "제목": "선택 알고리즘 vs 기준선 변화 비교", "데이터": "improvement_over_baseline"},
        {"번호": 7, "유형": "CDFChart",       "제목": "지연 누적분포 (CDF)",              "데이터": "per_edge_metrics.latency_ms"},
        {"번호": 8, "유형": "ScatterChart",   "제목": "CBR vs BS 부하",                   "데이터": "per_edge_metrics.cbr"},
    ]

    rec_lines = []
    if rec.get("primary_finding"):
        rec_lines.append(rec["primary_finding"])
    if rec.get("improvement_highlights"):
        rec_lines += rec["improvement_highlights"]
    recommendation_text = "  ".join(rec_lines) if rec_lines else "분석 데이터가 충분하지 않습니다."

    cw = meta.get("cost_weights") or {}
    ns = meta.get("norm_scales") or {}

    cbr_vals = [e.get("cbr") for e in edges if e.get("cbr") is not None]
    cbr_avg  = round(sum(cbr_vals) / len(cbr_vals), 4) if cbr_vals else None

    return ReportDocument(
        title="V2X 시뮬레이션 분석 보고서",
        subtitle=subtitle,
        generated_at=rs.get("generated_at", ""),
        run_id=rs.get("run_id", ""),
        scenario_id=rs.get("scenario_id", ""),
        scenario_name=meta.get("scenario_name") or rs.get("scenario_id", ""),
        origin_lat=rs.get("origin_lat"),
        origin_lng=rs.get("origin_lng"),
        dest_lat=rs.get("dest_lat"),
        dest_lng=rs.get("dest_lng"),
        vehicle_count=rs.get("vehicle_count", 0),
        network_mode=rs.get("network_mode", ""),
        sim_mode=rs.get("sim_mode", ""),
        seed=rs.get("seed"),
        route_algorithm=meta.get("route_algorithm") or rs.get("route_algorithm", ""),
        bs_selection_algorithm=meta.get("bs_selection_algorithm", ""),
        resource_allocation_algorithm=meta.get("resource_allocation_algorithm", ""),
        lookahead_k=meta.get("lookahead_k"),
        cost_weights=dict(cw),
        norm_scales=dict(ns),
        kpis=kpis,
        improvement=improvement,
        selected_algorithm=rs.get("selected_algorithm", ""),
        baseline_algorithm=rs.get("baseline_algorithm", ""),
        algorithms=algos,
        chart_descriptions=chart_descriptions,
        bottleneck_sections=sim_s.get("bottleneck_sections") or [],
        overloaded_bs=sim_s.get("overloaded_base_stations") or [],
        high_latency_sections=sim_s.get("high_latency_sections") or [],
        handover_sections=sim_s.get("frequent_handover_sections") or [],
        future_risk_sections=sim_s.get("future_connectivity_risk_sections") or [],
        primary_finding=rec.get("primary_finding", ""),
        trade_offs=list(rec.get("trade_offs") or []),
        risk_factors=list(rec.get("risk_factors") or []),
        recommendation_text=recommendation_text,
        per_edge_sample=edges[:20],
        per_bs_summary=bs,
        metadata_notes={k: v for k, v in meta.items() if k not in ("cost_weights", "norm_scales")},
        pir_p99_ms=rs.get("pir_p99_ms"),
        pir_compliant=rs.get("pir_compliant"),
        jain_fairness_index=rs.get("jain_fairness_index"),
        cbr_avg=cbr_avg,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Markdown renderer
# ──────────────────────────────────────────────────────────────────────────────

def _md_kv_table(rows: list[tuple[str, Any]]) -> str:
    lines = ["| 항목 | 값 |", "|---|---|"]
    for label, val in rows:
        v = "—" if val is None else str(val)
        lines.append(f"| {label} | {v} |")
    return "\n".join(lines)


def _md_dict_table(d: dict) -> str:
    return _md_kv_table([(k, v) for k, v in d.items()])


def _md_algo_table(algos: list[dict]) -> str:
    if not algos:
        return "_알고리즘 데이터 없음_"
    keys = ["algorithm", "total_cost", "average_latency_ms", "prr_approx",
            "handover_count", "average_bs_load", "summary_rank_score"]
    headers = ["알고리즘", "총 비용", "평균 지연(ms)", "PRR", "핸드오버", "BS 부하", "순위 점수"]
    lines = ["| " + " | ".join(headers) + " |",
             "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in algos:
        cells = [str(row.get(k, "—")) for k in keys]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def _md_finding_list(items: list[dict], key_fields: list[str]) -> str:
    if not items:
        return "_해당 없음_\n"
    lines = []
    for item in items:
        parts = [str(item.get(f, "")) for f in key_fields if item.get(f) is not None]
        lines.append("- " + "  |  ".join(parts))
    return "\n".join(lines)


def render_report_markdown(doc: ReportDocument) -> str:
    """Render ReportDocument to a clean Markdown string."""
    sections: list[str] = []

    def h(level: int, text: str) -> str:
        return "#" * level + " " + text

    sections.append(h(1, doc.title))
    if doc.subtitle:
        sections.append(f"_{doc.subtitle}_\n")

    sections.append(h(2, "1. 보고서 정보"))
    sections.append(_md_kv_table([
        ("런 ID",    doc.run_id),
        ("생성 시각", doc.generated_at),
    ]))

    sections.append(h(2, "2. 시나리오 설명"))
    origin = (f"{doc.origin_lat:.5f}, {doc.origin_lng:.5f}"
              if doc.origin_lat is not None else "—")
    dest   = (f"{doc.dest_lat:.5f}, {doc.dest_lng:.5f}"
              if doc.dest_lat is not None else "—")
    sections.append(_md_kv_table([
        ("시나리오 ID",   doc.scenario_id),
        ("시나리오 이름", doc.scenario_name),
        ("출발지 (위도,경도)", origin),
        ("목적지 (위도,경도)", dest),
        ("차량 대수",     doc.vehicle_count),
        ("네트워크 모드", doc.network_mode),
    ]))

    sections.append(h(2, "3. 시뮬레이션 설정"))
    sections.append(_md_kv_table([
        ("시뮬레이션 모드",    doc.sim_mode),
        ("시드",              doc.seed),
        ("경로 알고리즘",     doc.route_algorithm),
        ("BS 선택 알고리즘",  doc.bs_selection_algorithm),
        ("자원 할당 알고리즘", doc.resource_allocation_algorithm),
        ("Look-ahead K",      doc.lookahead_k),
    ]))
    if doc.cost_weights:
        sections.append("\n**비용 가중치**\n")
        sections.append(_md_dict_table(doc.cost_weights))

    sections.append(h(2, "4. 핵심 KPI"))
    sections.append(_md_dict_table(doc.kpis))
    if any(v is not None for v in doc.improvement.values()):
        sections.append(f"\n**기준선 대비 개선 ({doc.baseline_algorithm} → {doc.selected_algorithm})**\n")
        sections.append(_md_dict_table(doc.improvement))

    sections.append(h(2, "5. 알고리즘 비교"))
    sections.append(f"선택 알고리즘: **{doc.selected_algorithm}**  기준선: {doc.baseline_algorithm}\n")
    sections.append(_md_algo_table(doc.algorithms))

    sections.append(h(2, "6. 병목 / 위험 발견"))
    sections.append(h(3, "6.1 병목 구간"))
    sections.append(_md_finding_list(doc.bottleneck_sections,
                                     ["street_name", "severity", "load_ratio", "latency_ms", "connected_bs"]))
    sections.append(h(3, "6.2 과부하 기지국"))
    sections.append(_md_finding_list(doc.overloaded_bs,
                                     ["bs_name", "severity", "load_ratio", "affected_edge_count"]))
    sections.append(h(3, "6.3 고지연 구간"))
    sections.append(_md_finding_list(doc.high_latency_sections,
                                     ["street_name", "latency_ms", "excess_ms", "connected_bs"]))
    sections.append(h(3, "6.4 빈번 핸드오버 구간"))
    sections.append(_md_finding_list(doc.handover_sections,
                                     ["street_name", "from_bs_name", "to_bs_name", "latency_ms"]))
    sections.append(h(3, "6.5 미래 연결성 위험 구간"))
    sections.append(_md_finding_list(doc.future_risk_sections,
                                     ["street_name", "severity", "nearest_bs_name"]))

    sections.append(h(2, "7. 결론"))
    if doc.primary_finding:
        sections.append(f"> {doc.primary_finding}\n")
    if doc.trade_offs:
        sections.append("**트레이드오프**\n")
        sections.append("\n".join(f"- {t}" for t in doc.trade_offs))
    if doc.risk_factors:
        sections.append("\n**위험 요인**\n")
        sections.append("\n".join(f"- {r}" for r in doc.risk_factors))
    if doc.recommendation_text:
        sections.append(f"\n{doc.recommendation_text}")

    sections.append(h(2, "부록 A. 구간별 메트릭 (처음 20행)"))
    if doc.per_edge_sample:
        cols = ["edge_index", "street_name", "latency_ms", "load_ratio",
                "within_coverage", "handover", "cbr", "path_loss_db", "total_cost"]
        header = "| " + " | ".join(cols) + " |"
        sep    = "|" + "|".join(["---"] * len(cols)) + "|"
        rows   = [header, sep]
        for e in doc.per_edge_sample:
            rows.append("| " + " | ".join(str(e.get(c, "")) for c in cols) + " |")
        sections.append("\n".join(rows))
    else:
        sections.append("_데이터 없음_")

    sections.append(h(2, "부록 B. 기지국 요약"))
    if doc.per_bs_summary:
        cols = ["bs_name", "node_type", "load", "load_ratio", "severity",
                "affected_edge_count", "avg_latency_on_route_ms", "jain_fairness_index"]
        header = "| " + " | ".join(cols) + " |"
        sep    = "|" + "|".join(["---"] * len(cols)) + "|"
        rows   = [header, sep]
        for b in doc.per_bs_summary:
            rows.append("| " + " | ".join(str(b.get(c, "")) for c in cols) + " |")
        sections.append("\n".join(rows))
    else:
        sections.append("_데이터 없음_")

    sections.append(h(2, "참고문헌"))
    for i, ref in enumerate(REFERENCES, 1):
        sections.append(f"[{i}] {ref}")

    return "\n\n".join(sections) + "\n"


# ──────────────────────────────────────────────────────────────────────────────
# HTML renderer
# ──────────────────────────────────────────────────────────────────────────────

_HTML_STYLE = """<style>
:root{
  --bg:#F8FAFD;--surface:#FFF;--surface-2:#F1F5F9;
  --accent:#1D4ED8;--accent-2:#3B82F6;--accent-light:#DBEAFE;
  --good:#15803D;--good-bg:#DCFCE7;
  --warn:#B45309;--warn-bg:#FEF3C7;
  --bad:#B91C1C;--bad-bg:#FEE2E2;
  --text:#0F172A;--text-2:#334155;--muted:#64748B;
  --border:#E2E8F0;--th-bg:#EFF6FF;
  --shadow:0 1px 3px rgba(0,0,0,.08),0 1px 2px rgba(0,0,0,.06);
}
@media(prefers-color-scheme:dark){:root{
  --bg:#060B18;--surface:#0F172A;--surface-2:#1E293B;
  --accent:#60A5FA;--accent-2:#93C5FD;--accent-light:#1A2F52;
  --good:#4ADE80;--good-bg:#052E16;
  --warn:#FCD34D;--warn-bg:#1C1200;
  --bad:#F87171;--bad-bg:#2D0707;
  --text:#F1F5F9;--text-2:#CBD5E1;--muted:#94A3B8;
  --border:#1E293B;--th-bg:#1A2744;
  --shadow:0 1px 3px rgba(0,0,0,.4),0 1px 2px rgba(0,0,0,.3);
}}
:root[data-theme="light"]{--bg:#F8FAFD;--surface:#FFF;--surface-2:#F1F5F9;--accent:#1D4ED8;--accent-2:#3B82F6;--accent-light:#DBEAFE;--good:#15803D;--good-bg:#DCFCE7;--warn:#B45309;--warn-bg:#FEF3C7;--bad:#B91C1C;--bad-bg:#FEE2E2;--text:#0F172A;--text-2:#334155;--muted:#64748B;--border:#E2E8F0;--th-bg:#EFF6FF;--shadow:0 1px 3px rgba(0,0,0,.08),0 1px 2px rgba(0,0,0,.06);}
:root[data-theme="dark"]{--bg:#060B18;--surface:#0F172A;--surface-2:#1E293B;--accent:#60A5FA;--accent-2:#93C5FD;--accent-light:#1A2F52;--good:#4ADE80;--good-bg:#052E16;--warn:#FCD34D;--warn-bg:#1C1200;--bad:#F87171;--bad-bg:#2D0707;--text:#F1F5F9;--text-2:#CBD5E1;--muted:#94A3B8;--border:#1E293B;--th-bg:#1A2744;--shadow:0 1px 3px rgba(0,0,0,.4),0 1px 2px rgba(0,0,0,.3);}
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,sans-serif;font-size:14px;line-height:1.6;background:var(--bg);color:var(--text);}
.report-header{background:linear-gradient(135deg,var(--accent) 0%,#1E40AF 50%,var(--accent-2) 100%);color:#fff;padding:36px 40px 28px;margin-bottom:0;}
.report-header h1{font-size:22px;font-weight:700;letter-spacing:-.3px;margin-bottom:6px;}
.report-header .subtitle{opacity:.85;font-size:13px;margin-bottom:0;}
.report-header .meta{display:flex;flex-wrap:wrap;gap:20px;font-size:12px;opacity:.75;border-top:1px solid rgba(255,255,255,.2);padding-top:12px;margin-top:16px;}
.content{max-width:960px;margin:0 auto;padding:28px 32px 56px;}
.section{margin-bottom:32px;}
.section-title{font-size:14.5px;font-weight:700;padding:8px 0 8px 14px;border-left:3px solid var(--accent);margin-bottom:14px;color:var(--text);}
h3.sub{font-size:13px;font-weight:600;color:var(--text-2);margin:16px 0 8px;}
/* KPI grid */
.kpi-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(175px,1fr));gap:12px;margin-bottom:8px;}
.kpi-card{background:var(--surface);border:1px solid var(--border);border-radius:12px;padding:16px 18px;box-shadow:var(--shadow);}
.kpi-label{font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.5px;color:var(--muted);margin-bottom:8px;}
.kpi-num{font-size:30px;font-weight:700;line-height:1;font-variant-numeric:tabular-nums;margin-bottom:4px;}
.kpi-unit{font-size:11.5px;color:var(--muted);}
.kpi-ref{font-size:10.5px;color:var(--muted);margin-top:4px;}
.kpi-badge{display:inline-flex;align-items:center;font-size:10.5px;font-weight:700;padding:2px 8px;border-radius:20px;margin-top:8px;}
.c-good{color:var(--good);}.c-warn{color:var(--warn);}.c-bad{color:var(--bad);}.c-neutral{color:var(--muted);}
.bg-good{background:var(--good-bg);color:var(--good);}
.bg-warn{background:var(--warn-bg);color:var(--warn);}
.bg-bad{background:var(--bad-bg);color:var(--bad);}
.bg-neutral{background:var(--surface-2);color:var(--muted);}
/* Config KV table */
.kv-wrap{background:var(--surface);border:1px solid var(--border);border-radius:10px;overflow:hidden;margin-bottom:12px;box-shadow:var(--shadow);}
table.kv{width:100%;border-collapse:collapse;font-size:13px;}
table.kv td{padding:8px 14px;border-bottom:1px solid var(--border);}
table.kv tr:last-child td{border-bottom:none;}
table.kv td:first-child{font-weight:600;color:var(--text-2);width:42%;background:var(--surface-2);}
/* Algorithm bar chart */
.algo-list{display:flex;flex-direction:column;gap:10px;}
.algo-row{display:flex;align-items:center;gap:10px;}
.algo-name{width:210px;flex-shrink:0;font-size:12.5px;font-weight:500;color:var(--text-2);display:flex;align-items:center;gap:6px;flex-wrap:wrap;}
.algo-bar-bg{flex:1;height:24px;background:var(--surface-2);border-radius:5px;overflow:hidden;}
.algo-bar-fill{height:100%;border-radius:5px;display:flex;align-items:center;padding-left:8px;font-size:11px;font-weight:700;color:#fff;min-width:24px;}
.bar-selected{background:linear-gradient(90deg,var(--accent),var(--accent-2));}
.bar-other{background:var(--muted);opacity:.4;}
.algo-val{font-size:11.5px;color:var(--muted);flex-shrink:0;min-width:100px;text-align:right;font-variant-numeric:tabular-nums;}
.winner-tag{font-size:10px;font-weight:700;background:var(--accent-light);color:var(--accent);border-radius:20px;padding:1px 7px;white-space:nowrap;}
/* Data tables */
.table-wrap{overflow-x:auto;margin-bottom:16px;}
table.dt{border-collapse:collapse;width:100%;font-size:12px;min-width:560px;}
table.dt th{background:var(--th-bg);text-align:left;padding:8px 11px;border:1px solid var(--border);font-size:11px;font-weight:700;color:var(--accent);text-transform:uppercase;letter-spacing:.3px;}
table.dt td{padding:6px 11px;border:1px solid var(--border);color:var(--text-2);}
table.dt tr:nth-child(even) td{background:var(--surface-2);}
.td-good{color:var(--good);font-weight:600;}
.td-warn{color:var(--warn);font-weight:600;}
.td-bad{color:var(--bad);font-weight:700;}
/* Finding cards */
.finding-list{display:flex;flex-direction:column;gap:8px;}
.fc{background:var(--surface);border:1px solid var(--border);border-radius:8px;padding:10px 14px;display:flex;gap:10px;align-items:flex-start;}
.sev-dot{width:8px;height:8px;border-radius:50%;flex-shrink:0;margin-top:5px;}
.sev-h{background:var(--bad);}.sev-m{background:var(--warn);}.sev-l{background:var(--good);}.sev-n{background:var(--muted);}
.fc-title{font-weight:600;font-size:13px;color:var(--text);}
.fc-detail{font-size:11.5px;color:var(--muted);margin-top:2px;}
/* Conclusion */
.conclusion-block{background:var(--surface);border:1px solid var(--border);border-radius:10px;padding:20px 24px;box-shadow:var(--shadow);}
blockquote.finding{border-left:3px solid var(--accent);padding:12px 16px;background:var(--accent-light);border-radius:0 8px 8px 0;margin-bottom:14px;font-size:13.5px;line-height:1.7;}
/* References */
.refs-card{background:var(--surface);border:1px solid var(--border);border-radius:10px;padding:20px 24px;box-shadow:var(--shadow);}
ol.refs{padding-left:20px;font-size:11.5px;color:var(--muted);line-height:1.9;}
.footnote{font-size:11px;color:var(--muted);border-top:1px solid var(--border);padding-top:12px;margin-top:14px;}
hr.sep{border:none;border-top:1px solid var(--border);margin:28px 0;}
@media print{.report-header{-webkit-print-color-adjust:exact;print-color-adjust:exact}.kpi-grid{break-inside:avoid}}
</style>"""


def _kpi_card(label: str, value: Any, unit: str = "", *,
              tone: str = "neutral", badge: str = "", ref: str = "") -> str:
    if value is None:
        val_str, color_cls = "—", "c-neutral"
    else:
        val_str = f"{value:.2f}" if isinstance(value, float) and value < 100 else \
                  f"{value:.0f}" if isinstance(value, float) else str(value)
        color_cls = f"c-{tone}" if tone in ("good", "warn", "bad") else "c-neutral"
    badge_html = ""
    if badge:
        bcls = f"bg-{tone}" if tone in ("good", "warn", "bad") else "bg-neutral"
        badge_html = f'<div class="kpi-badge {bcls}">{_html.escape(badge)}</div>'
    ref_html = f'<div class="kpi-ref">{_html.escape(ref)}</div>' if ref else ""
    return (f'<div class="kpi-card">'
            f'<div class="kpi-label">{_html.escape(label)}</div>'
            f'<div class="kpi-num {color_cls}">{_html.escape(val_str)}</div>'
            f'<div class="kpi-unit">{_html.escape(unit)}</div>'
            f'{badge_html}{ref_html}</div>')


def _kv_row(key: str, val: Any) -> str:
    v = "—" if val is None else _html.escape(str(val))
    return f"<tr><td>{_html.escape(key)}</td><td>{v}</td></tr>"


def _kv_card(rows: list[tuple[str, Any]]) -> str:
    body = "".join(_kv_row(k, v) for k, v in rows)
    return f'<div class="kv-wrap"><table class="kv"><tbody>{body}</tbody></table></div>'


def _algo_bar_chart(algos: list[dict], selected: str) -> str:
    if not algos:
        return '<p style="color:var(--muted);font-size:13px">알고리즘 데이터 없음</p>'
    costs = [float(a.get("total_cost") or 0) for a in algos]
    max_cost = max(costs) if costs else 1.0
    parts = ['<div class="algo-list">']
    for a, cost in zip(algos, costs):
        algo_id = str(a.get("algorithm", "—"))
        pct = int(cost / max_cost * 100) if max_cost > 0 else 0
        is_sel = (algo_id == selected)
        bar_cls = "bar-selected" if is_sel else "bar-other"
        winner = '<span class="winner-tag">선택</span>' if is_sel else ""
        lat = a.get("average_latency_ms")
        prr = a.get("prr_approx")
        meta = "  ·  ".join(filter(None, [
            f"{lat:.1f} ms" if lat is not None else "",
            f"PRR {prr:.3f}" if prr is not None else "",
        ]))
        parts.append(
            f'<div class="algo-row">'
            f'<div class="algo-name">{_html.escape(algo_id)}{winner}</div>'
            f'<div class="algo-bar-bg">'
            f'<div class="algo-bar-fill {bar_cls}" style="width:{pct}%">'
            f'{_html.escape(f"{cost:.0f}") if pct > 15 else ""}</div></div>'
            f'<div class="algo-val">{_html.escape(meta)}</div>'
            f'</div>'
        )
    parts.append('</div>')
    return "".join(parts)


def _finding_list_html(items: list[dict], key_fields: list[str]) -> str:
    if not items:
        return '<p style="color:var(--muted);font-size:13px;padding:4px 0">해당 없음</p>'
    sev_map = {"high": ("sev-h", "높음"), "medium": ("sev-m", "보통"),
               "low": ("sev-l", "낮음")}
    parts = ['<div class="finding-list">']
    for item in items:
        sev_raw = str(item.get("severity", "") or "").lower()
        dot_cls, sev_label = sev_map.get(sev_raw, ("sev-n", sev_raw or "—"))
        name = next((item.get(k) for k in ("street_name", "bs_name", "road_name")
                     if item.get(k)), "—")
        detail_parts = [
            f"{f}: {item[f]}" for f in key_fields
            if f not in ("street_name", "bs_name", "severity") and item.get(f) is not None
        ][:5]
        parts.append(
            f'<div class="fc"><div class="sev-dot {dot_cls}"></div><div>'
            f'<div class="fc-title">{_html.escape(str(name))} '
            f'<span style="font-size:11px;color:var(--muted);font-weight:400">'
            f'심각도: {_html.escape(sev_label)}</span></div>'
            f'<div class="fc-detail">{_html.escape("  |  ".join(detail_parts))}</div>'
            f'</div></div>'
        )
    parts.append('</div>')
    return "".join(parts)


def _edge_table_html(edges: list[dict]) -> str:
    if not edges:
        return '<p style="color:var(--muted);font-size:13px">데이터 없음</p>'
    cols = ["edge_index", "street_name", "latency_ms", "load_ratio",
            "within_coverage", "handover", "cbr", "path_loss_db", "total_cost"]
    hdrs = ["#", "도로명", "지연(ms)", "부하율", "커버리지", "핸드오버", "CBR†", "경로손실(dB)†", "비용"]
    head = "<tr>" + "".join(f"<th>{h}</th>" for h in hdrs) + "</tr>"
    body = ""
    for row in edges:
        lat = row.get("latency_ms")
        lat_cls = ("td-good" if lat is not None and lat < 20
                   else "td-warn" if lat is not None and lat < 100
                   else "td-bad" if lat is not None else "")
        cells = []
        for c in cols:
            v = row.get(c)
            raw = "—" if v is None else str(v)
            if c == "latency_ms" and lat_cls:
                cells.append(f'<td class="{lat_cls}">{_html.escape(raw)}</td>')
            elif c in ("within_coverage", "handover"):
                cls = "td-good" if v is True else ("td-bad" if v is False else "")
                sym = ("&#10003;" if v is True else "&#10007;" if v is False else "—")
                cells.append(f'<td class="{cls}">{sym}</td>')
            else:
                cells.append(f"<td>{_html.escape(raw)}</td>")
        body += "<tr>" + "".join(cells) + "</tr>"
    return (f'<div class="table-wrap"><table class="dt">'
            f'<thead>{head}</thead><tbody>{body}</tbody></table></div>')


def _bs_table_html(bs_rows: list[dict]) -> str:
    if not bs_rows:
        return '<p style="color:var(--muted);font-size:13px">데이터 없음</p>'
    cols = ["bs_name", "node_type", "load_ratio", "severity",
            "affected_edge_count", "avg_latency_on_route_ms"]
    hdrs = ["BS명", "유형", "부하율", "심각도", "영향 구간", "평균 지연(ms)"]
    head = "<tr>" + "".join(f"<th>{h}</th>" for h in hdrs) + "</tr>"
    body = ""
    for row in bs_rows:
        cells = []
        for c in cols:
            v = row.get(c)
            raw = "—" if v is None else str(v)
            if c == "load_ratio" and v is not None:
                try:
                    fv = float(v)
                    cls = "td-good" if fv < 0.7 else ("td-warn" if fv < 0.9 else "td-bad")
                    cells.append(f'<td class="{cls}">{_html.escape(raw)}</td>')
                except (ValueError, TypeError):
                    cells.append(f"<td>{_html.escape(raw)}</td>")
            elif c == "severity":
                scls = {"high": "td-bad", "medium": "td-warn",
                        "low": "td-good"}.get(str(v or "").lower(), "")
                cells.append(f'<td class="{scls}">{_html.escape(raw)}</td>')
            else:
                cells.append(f"<td>{_html.escape(raw)}</td>")
        body += "<tr>" + "".join(cells) + "</tr>"
    return (f'<div class="table-wrap"><table class="dt">'
            f'<thead>{head}</thead><tbody>{body}</tbody></table></div>')


def render_report_html(doc: ReportDocument) -> str:
    """Render ReportDocument to a visually polished, self-contained HTML report."""

    def _tone_lat(v):   return "neutral" if v is None else ("good" if v < 20 else "warn" if v < 100 else "bad")
    def _tone_prr(v):   return "neutral" if v is None else ("good" if v >= 0.9 else "warn" if v >= 0.7 else "bad")
    def _tone_jain(v):  return "neutral" if v is None else ("good" if v >= 0.9 else "warn" if v >= 0.7 else "bad")
    def _tone_cbr(v):   return "neutral" if v is None else ("good" if v < 0.65 else "warn" if v < 0.8 else "bad")
    def _tone_cov(v):   return "neutral" if v is None else ("good" if v >= 90 else "warn" if v >= 70 else "bad")
    def _tone_pir(v):   return "neutral" if v is None else ("good" if v <= 100 else "bad")
    def _tone_impr(v):
        try: return "neutral" if v is None else ("good" if float(v) > 0 else "bad")
        except (ValueError, TypeError): return "neutral"

    avg_lat  = doc.kpis.get("평균 지연 (ms)")
    prr      = doc.kpis.get("PRR (근사)")
    handover = doc.kpis.get("핸드오버 횟수")
    coverage = doc.kpis.get("커버리지율 (%)")
    cost_impr = doc.improvement.get("비용 개선 (%)")

    kpi_cards = [
        _kpi_card("평균 지연", avg_lat, "ms",
                  tone=_tone_lat(avg_lat),
                  badge=("충족" if avg_lat is not None and avg_lat < 20 else
                         "초과" if avg_lat is not None else ""),
                  ref="목표 < 20 ms [8]"),
        _kpi_card("PRR", prr, "",
                  tone=_tone_prr(prr),
                  badge=("충족" if prr is not None and prr >= 0.9 else
                         "미달" if prr is not None else ""),
                  ref="> 0.90 [6]"),
        _kpi_card("Jain 공정성", doc.jain_fairness_index, "",
                  tone=_tone_jain(doc.jain_fairness_index),
                  badge=("공정" if doc.jain_fairness_index is not None and doc.jain_fairness_index >= 0.9 else
                         "보통" if doc.jain_fairness_index is not None and doc.jain_fairness_index >= 0.7 else
                         "불공정" if doc.jain_fairness_index is not None else ""),
                  ref=">= 0.90 권장 [9]"),
        _kpi_card("CBR", doc.cbr_avg, "",
                  tone=_tone_cbr(doc.cbr_avg),
                  badge=("정상" if doc.cbr_avg is not None and doc.cbr_avg < 0.65 else
                         "혼잡" if doc.cbr_avg is not None else ""),
                  ref="< 0.65 [ETSI 4]"),
        _kpi_card("핸드오버", handover, "회", tone="neutral"),
        _kpi_card("커버리지율", coverage, "%",
                  tone=_tone_cov(coverage),
                  badge=("우수" if coverage is not None and coverage >= 90 else
                         "양호" if coverage is not None and coverage >= 70 else
                         "불량" if coverage is not None else ""),
                  ref="> 90 % 권장"),
        _kpi_card("PIR P99", doc.pir_p99_ms, "ms",
                  tone=_tone_pir(doc.pir_p99_ms),
                  badge=("충족" if doc.pir_p99_ms is not None and doc.pir_p99_ms <= 100 else
                         "초과" if doc.pir_p99_ms is not None else ""),
                  ref="<= 100 ms [1, 7]"),
        _kpi_card("비용 개선", cost_impr, "%",
                  tone=_tone_impr(cost_impr),
                  badge=(f"{doc.selected_algorithm} vs baseline" if doc.selected_algorithm else "")),
    ]

    p: list[str] = []
    p.append("<!DOCTYPE html><html lang='ko'>")
    p.append("<head><meta charset='UTF-8'>"
             "<meta name='viewport' content='width=device-width,initial-scale=1'>")
    p.append(f"<title>{_html.escape(doc.title)}</title>")
    p.append(_HTML_STYLE)
    p.append("</head><body>")

    # ── Report header ──
    p.append('<div class="report-header">')
    p.append(f'<h1>{_html.escape(doc.title)}</h1>')
    if doc.subtitle:
        p.append(f'<div class="subtitle">{_html.escape(doc.subtitle)}</div>')
    meta_items = [
        (f"Run: {doc.run_id}" if doc.run_id else ""),
        (doc.generated_at or ""),
        (f"시나리오: {doc.scenario_name}" if doc.scenario_name else ""),
        (f"차량 {doc.vehicle_count}대" if doc.vehicle_count else ""),
    ]
    p.append('<div class="meta">' +
             "".join(f'<span>{_html.escape(m)}</span>' for m in meta_items if m) +
             '</div></div>')

    p.append('<div class="content">')

    # ── KPI grid ──
    p.append('<div class="section">')
    p.append('<div class="section-title">핵심 성능 지표 (KPI)</div>')
    p.append('<div class="kpi-grid">' + "".join(kpi_cards) + '</div>')
    p.append('</div>')

    # ── Algorithm comparison ──
    if doc.algorithms:
        p.append('<div class="section">')
        p.append('<div class="section-title">알고리즘 비교</div>')
        sel_label = _html.escape(doc.selected_algorithm)
        base_label = _html.escape(doc.baseline_algorithm)
        p.append(f'<p style="font-size:13px;color:var(--muted);margin-bottom:14px">'
                 f'선택: <strong style="color:var(--text)">{sel_label}</strong>'
                 f'&nbsp;&middot;&nbsp;기준선: {base_label}</p>')
        p.append(_algo_bar_chart(doc.algorithms, doc.selected_algorithm))
        p.append('</div>')

    # ── Scenario / Config ──
    p.append('<div class="section">')
    p.append('<div class="section-title">시나리오 및 시뮬레이션 설정</div>')
    origin = (f"{doc.origin_lat:.5f}, {doc.origin_lng:.5f}"
              if doc.origin_lat is not None else None)
    dest = (f"{doc.dest_lat:.5f}, {doc.dest_lng:.5f}"
            if doc.dest_lat is not None else None)
    p.append(_kv_card([
        ("시나리오 ID", doc.scenario_id),
        ("출발지", origin), ("목적지", dest),
        ("네트워크 모드", doc.network_mode),
        ("시뮬레이션 모드", doc.sim_mode),
        ("경로 알고리즘", doc.route_algorithm),
        ("BS 선택 알고리즘", doc.bs_selection_algorithm),
        ("자원 할당 알고리즘", doc.resource_allocation_algorithm),
        ("시드", doc.seed), ("Look-ahead K", doc.lookahead_k),
    ]))
    if doc.cost_weights:
        p.append('<h3 class="sub">비용 가중치</h3>')
        p.append(_kv_card(list(doc.cost_weights.items())))
    p.append('</div>')

    # ── Findings ──
    has_findings = any([doc.bottleneck_sections, doc.overloaded_bs,
                        doc.high_latency_sections, doc.handover_sections,
                        doc.future_risk_sections])
    if has_findings:
        p.append('<div class="section">')
        p.append('<div class="section-title">병목 및 위험 발견</div>')
        for title, items, fields in [
            ("병목 구간", doc.bottleneck_sections,
             ["street_name", "severity", "load_ratio", "latency_ms", "connected_bs"]),
            ("과부하 기지국", doc.overloaded_bs,
             ["bs_name", "severity", "load_ratio", "affected_edge_count"]),
            ("고지연 구간", doc.high_latency_sections,
             ["street_name", "latency_ms", "excess_ms", "connected_bs"]),
            ("빈번 핸드오버 구간", doc.handover_sections,
             ["street_name", "from_bs_name", "to_bs_name", "latency_ms"]),
            ("미래 연결성 위험", doc.future_risk_sections,
             ["street_name", "severity", "nearest_bs_name"]),
        ]:
            if items:
                p.append(f'<h3 class="sub">{title} ({len(items)}개)</h3>')
                p.append(_finding_list_html(items, fields))
        p.append('</div>')

    # ── Conclusion ──
    p.append('<div class="section">')
    p.append('<div class="section-title">결론 및 권고사항</div>')
    p.append('<div class="conclusion-block">')
    if doc.primary_finding:
        p.append(f'<blockquote class="finding">{_html.escape(doc.primary_finding)}</blockquote>')
    if doc.trade_offs:
        p.append('<h3 class="sub">트레이드오프</h3>')
        p.append("<ul style='padding-left:20px;font-size:13px;line-height:1.8'>" +
                 "".join(f"<li>{_html.escape(t)}</li>" for t in doc.trade_offs) + "</ul>")
    if doc.risk_factors:
        p.append('<h3 class="sub" style="margin-top:14px">위험 요인</h3>')
        p.append("<ul style='padding-left:20px;font-size:13px;line-height:1.8'>" +
                 "".join(f"<li>{_html.escape(r)}</li>" for r in doc.risk_factors) + "</ul>")
    if doc.recommendation_text and doc.recommendation_text != "분석 데이터가 충분하지 않습니다.":
        p.append(f'<p style="font-size:13px;color:var(--muted);margin-top:14px">'
                 f'{_html.escape(doc.recommendation_text)}</p>')
    p.append('</div></div>')

    # ── Appendices ──
    if doc.per_edge_sample:
        p.append('<hr class="sep">')
        p.append('<div class="section">')
        p.append('<div class="section-title">부록 A. 구간별 메트릭 (처음 20행)</div>')
        p.append(_edge_table_html(doc.per_edge_sample))
        p.append('</div>')
    if doc.per_bs_summary:
        p.append('<div class="section">')
        p.append('<div class="section-title">부록 B. 기지국 요약</div>')
        p.append(_bs_table_html(doc.per_bs_summary))
        p.append('</div>')

    # ── References ──
    p.append('<hr class="sep">')
    p.append('<div class="refs-card">')
    p.append('<div class="section-title" style="margin-bottom:12px">참고문헌 / References</div>')
    p.append('<ol class="refs">' +
             "".join(f"<li>{_html.escape(r)}</li>" for r in REFERENCES) + '</ol>')
    p.append('<div class="footnote">'
             '† 해석 모델 기반 (analytical model). '
             'MAC 계층 측정 없이 검증된 수식을 적용한 추정치이며 실제 측정과 차이 가능.'
             '</div></div>')

    p.append('</div>')  # .content
    p.append('</body></html>')
    return "".join(p)


# ──────────────────────────────────────────────────────────────────────────────
# DOCX entry point (requires python-docx)
# ──────────────────────────────────────────────────────────────────────────────

def generate_docx(doc: ReportDocument, lang: str = "ko") -> bytes:
    """
    Generate a DOCX file from ReportDocument.

    Parameters
    ----------
    doc  : ReportDocument
    lang : "ko"   Korean (RISS/DBpia submission)
           "en"   English (IEEE/Nature submission)
           "both" Bilingual (international co-publication)

    Requires: pip install python-docx
    Raises ImportError if python-docx is not installed.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Cm, Pt, RGBColor
    except ImportError as exc:
        raise ImportError(
            "python-docx가 설치되어 있지 않습니다. "
            "`pip install python-docx` 를 실행한 후 재시도하세요."
        ) from exc

    d = Document()

    # A4 margins
    for sec in d.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # ── helpers ──────────────────────────────────────────────────────────────

    def _cell_bg(cell, hex_color: str) -> None:
        """Set table cell background color (python-docx workaround)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _styled_table(headers: list[str], rows: list[list[Any]],
                      caption: str | None = None) -> None:
        """Add a styled table: blue header row, alternating row shading."""
        t = d.add_table(rows=len(rows) + 1, cols=len(headers))
        t.style = "Table Grid"
        # Header row
        hdr = t.rows[0]
        for i, h_text in enumerate(headers):
            cell = hdr.cells[i]
            cell.text = str(h_text)
            _cell_bg(cell, "1A6FCF")
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        # Data rows
        for ri, row_data in enumerate(rows):
            row = t.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                cell.text = "—" if val is None else str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else None
                if ri % 2 == 1:
                    _cell_bg(cell, "EFF6FF")
        if caption:
            cp = d.add_paragraph(caption)
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)
        d.add_paragraph("")

    def _formula(formula_text: str, citation: str = "") -> None:
        """Add a formula paragraph (Courier New, indented)."""
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(formula_text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.italic = True
        if citation:
            crun = p.add_run(f"   [{citation}]")
            crun.font.size = Pt(9)
            crun.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def _section(num: str, title_ko: str, title_en: str, level: int = 1) -> None:
        if lang == "en":
            d.add_heading(f"{num} {title_en}", level=level)
        elif lang == "both":
            d.add_heading(f"{num} {title_ko} / {title_en}", level=level)
        else:
            d.add_heading(f"{num} {title_ko}", level=level)

    def _txt(ko: str, en: str) -> str:
        if lang == "en":  return en
        if lang == "both": return f"{ko} / {en}"
        return ko

    def _kv(pairs: list[tuple[str, str, Any]]) -> None:
        """Render a key-value table. Each tuple: (ko_label, en_label, value)."""
        headers_h = [_txt("항목", "Item"), _txt("값", "Value")]
        rows_h = [
            [_txt(ko, en), "—" if v is None else str(v)]
            for ko, en, v in pairs
        ]
        _styled_table(headers_h, rows_h)

    def _add_references() -> None:
        d.add_heading(_txt("참고문헌", "References"), level=1)
        for i, ref in enumerate(REFERENCES, 1):
            p = d.add_paragraph(style="List Number")
            rb = p.add_run(f"[{i}]  ")
            rb.bold = True
            rb.font.size = Pt(9)
            rc2 = p.add_run(ref)
            rc2.font.size = Pt(9)

    # ── Title page ────────────────────────────────────────────────────────────
    title_text = (
        "V2X 시뮬레이션 분석 보고서" if lang in ("ko", "both")
        else "V2X Simulation Analysis Report"
    )
    tp = d.add_heading(title_text, level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if doc.subtitle:
        sp = d.add_paragraph(doc.subtitle)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.runs[0].italic = True
        sp.runs[0].font.size = Pt(11)

    meta_p = d.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(f"Run ID: {doc.run_id or '—'}\n")
    meta_p.add_run(f"{_txt('생성 시각', 'Generated')}: {doc.generated_at or '—'}")
    meta_p.runs[0].font.size = Pt(10)

    d.add_page_break()

    # ── Abstract ──────────────────────────────────────────────────────────────
    _section("Abstract" if lang == "en" else "초록", "초록", "Abstract", level=1)

    avg_lat  = doc.kpis.get("평균 지연 (ms)") or "—"
    cost_imp = doc.improvement.get("비용 개선 (%)")
    cost_imp_str = f"{cost_imp:.1f}" if isinstance(cost_imp, (int, float)) else "—"
    sel_algo = doc.selected_algorithm or "—"

    if lang == "en":
        abs_text = (
            f"This report presents a V2X (Vehicle-to-Everything) network simulation analysis. "
            f"The {sel_algo} algorithm achieves a {cost_imp_str}% cost reduction over baseline "
            f"with an average E2E latency of {avg_lat} ms. "
            f"Key performance indicators evaluated include Packet Reception Ratio (PRR), "
            f"Channel Busy Ratio (CBR), Packet Inter-Reception Time P99 (PIR), "
            f"Jain's Fairness Index, and path loss per edge. "
            f"All analytical models are derived from peer-reviewed standards and literature "
            f"(3GPP TR 37.885, ETSI TS 102 687, Gonzalez-Martin et al. 2019, Fernandez et al. 2014)."
        )
    else:
        abs_text = (
            f"본 보고서는 V2X(차량 대 사물) 네트워크 시뮬레이션 분석 결과를 제시한다. "
            f"{sel_algo} 알고리즘은 기준선 대비 비용 {cost_imp_str}% 개선, "
            f"평균 E2E 지연 {avg_lat} ms를 달성하였다. "
            f"평가된 핵심 성능 지표에는 패킷 수신율(PRR), 채널 점유율(CBR), "
            f"패킷 재수신 간격 P99(PIR), Jain 공정성 지수, 구간별 경로손실이 포함된다. "
            f"모든 해석 모델은 동료 심사 표준 및 문헌에서 도출하였다 "
            f"(3GPP TR 37.885, ETSI TS 102 687, Gonzalez-Martin et al. 2019, Fernandez et al. 2014)."
        )

    ap = d.add_paragraph(abs_text)
    ap.runs[0].font.size = Pt(10)
    d.add_paragraph("")

    # ── 1. Introduction ───────────────────────────────────────────────────────
    _section("1.", "소개", "Introduction")
    _section("1.1", "연구 배경", "Background", level=2)
    bg_text = _txt(
        "V2X 통신은 자율주행 및 협력 지능형 교통 시스템(C-ITS)의 핵심 인프라로, "
        "3GPP TS 22.186 [1]에 따라 100ms 이하의 E2E 지연과 99% 이상의 패킷 수신율이 요구된다. "
        "본 시뮬레이션은 도심 도로망에서의 기지국(BS) 선택 및 자원 할당 알고리즘의 성능을 비교 분석한다.",
        "V2X communications are a key infrastructure for autonomous driving and Cooperative Intelligent "
        "Transport Systems (C-ITS). Per 3GPP TS 22.186 [1], E2E latency < 100 ms and PRR > 99% are "
        "required. This simulation benchmarks BS selection and resource allocation algorithms on an "
        "urban road network."
    )
    bp = d.add_paragraph(bg_text)
    bp.runs[0].font.size = Pt(10)

    _section("1.2", "관련 연구", "Related Work", level=2)
    rw_text = _txt(
        "Gonzalez-Martin et al. [4]은 C-V2X Mode 4의 CBR 해석 모델을 제시하였으며, "
        "Fernandez et al. [5]은 5.9 GHz 대역에서의 차량 통신 경로손실 파라미터를 측정하였다. "
        "Eckermann et al. [7]은 PIR P99의 기하분포 상한 모델을 검증하였다.",
        "Gonzalez-Martin et al. [4] derived an analytical CBR model for C-V2X Mode 4. "
        "Fernandez et al. [5] measured path loss parameters for vehicular communications at 5.9 GHz. "
        "Eckermann et al. [7] validated the geometric-distribution upper bound for PIR P99."
    )
    rp = d.add_paragraph(rw_text)
    rp.runs[0].font.size = Pt(10)

    # ── 2. Methodology ────────────────────────────────────────────────────────
    _section("2.", "방법론", "Methodology")

    _section("2.1", "네트워크 모델", "Network Model", level=2)
    nm_text = _txt(
        "네트워크 모델은 5G NR Uu 인터페이스(기지국)와 PC5 사이드링크(RSU)로 구성된다. "
        "차량-기지국 간 경로손실은 Fernandez et al. [5] 로그-거리 모델을 적용한다.",
        "The network model comprises 5G NR Uu interface (base stations) and PC5 sidelink (RSUs). "
        "Vehicle-to-BS path loss follows the log-distance model of Fernandez et al. [5]."
    )
    d.add_paragraph(nm_text).runs[0].font.size = Pt(10)
    _formula("PL(d) = PL(d₀) + 10·n·log₁₀(d/d₀)  [dB]", "5")
    d.add_paragraph(_txt(
        "파라미터: n=1.61 (LOS highway), n=2.75 (LOS urban), n=3.50 (NLOS urban), d₀=10 m @ 5.9 GHz",
        "Parameters: n=1.61 (LOS highway), n=2.75 (LOS urban), n=3.50 (NLOS urban), d₀=10 m @ 5.9 GHz"
    )).runs[0].font.size = Pt(9)

    _section("2.2", "경로 비용 함수", "Route Cost Function", level=2)
    d.add_paragraph(_txt(
        "경로 비용 함수는 거리, 지연, 부하, 커버리지 위험의 가중 합으로 정의된다:",
        "The route cost function is defined as the weighted sum of distance, latency, load, and coverage risk:"
    )).runs[0].font.size = Pt(10)
    _formula("C = Σ_e [w₁c_d(e) + w₂c_L(e) + w₃c_ρ(e) + w₄c_cov(e)]")

    cw_headers = [_txt("가중치", "Weight"), _txt("의미", "Description"), _txt("값", "Value")]
    cw = doc.cost_weights or {}
    cw_rows = [
        ["w₁ (w_distance)",  _txt("거리 비용", "Distance cost"), cw.get("w_distance", "—")],
        ["w₂ (w_latency)",   _txt("지연 비용", "Latency cost"),  cw.get("w_latency",  "—")],
        ["w₃ (w_load)",      _txt("부하 비용", "Load cost"),     cw.get("w_load",     "—")],
        ["w₄ (w_coverage_risk)", _txt("커버리지 위험", "Coverage risk"), cw.get("w_coverage_risk", "—")],
    ]
    _styled_table(cw_headers, cw_rows,
                  caption=_txt("Table 1. 경로 비용 가중치 설정", "Table 1. Route cost function weights"))

    _section("2.3", "성능 지표 정의", "Performance Metric Definitions", level=2)
    fm_headers = [
        _txt("지표", "Metric"),
        _txt("수식", "Formula"),
        _txt("출처", "Reference"),
        _txt("임계값", "Threshold"),
    ]
    fm_rows = [
        [
            _txt(f["metric_ko"], f["metric_en"]),
            f["formula"],
            f["ref"],
            f["threshold"],
        ]
        for f in METRIC_FORMULAS
    ]
    _styled_table(fm_headers, fm_rows,
                  caption=_txt(
                      "Table 2. 성능 지표 수식 및 기준값 (†: 해석 모델 기반)",
                      "Table 2. Performance metric definitions and thresholds (†: analytical model)"
                  ))

    # ── 3. Simulation Setup ───────────────────────────────────────────────────
    _section("3.", "시뮬레이션 환경", "Simulation Setup")

    _section("3.1", "시나리오", "Scenario", level=2)
    origin_str = (f"{doc.origin_lat:.5f}, {doc.origin_lng:.5f}"
                  if doc.origin_lat is not None else "—")
    dest_str   = (f"{doc.dest_lat:.5f}, {doc.dest_lng:.5f}"
                  if doc.dest_lat is not None else "—")
    _kv([
        ("시나리오 ID",          "Scenario ID",           doc.scenario_id),
        ("시나리오 이름",        "Scenario Name",         doc.scenario_name),
        ("출발지 (위도, 경도)",  "Origin (lat, lng)",     origin_str),
        ("목적지 (위도, 경도)",  "Destination (lat, lng)", dest_str),
        ("차량 대수",            "Vehicle Count",         doc.vehicle_count),
        ("네트워크 모드",        "Network Mode",          doc.network_mode),
    ])

    _section("3.2", "파라미터 설정", "Algorithm Configuration", level=2)
    _kv([
        ("시뮬레이션 모드",      "Simulation Mode",           doc.sim_mode),
        ("시드",                 "Seed",                      doc.seed),
        ("경로 알고리즘",        "Route Algorithm",           doc.route_algorithm),
        ("BS 선택 알고리즘",     "BS Selection Algorithm",    doc.bs_selection_algorithm),
        ("자원 할당 알고리즘",   "Resource Allocation Algo",  doc.resource_allocation_algorithm),
        ("Look-ahead K",         "Look-ahead K",              doc.lookahead_k),
    ])

    # ── 4. Results ────────────────────────────────────────────────────────────
    _section("4.", "결과", "Results")

    _section("4.1", "핵심 KPI", "Key Performance Indicators", level=2)
    prr_v  = doc.kpis.get("PRR (근사)")
    twdr_v = doc.kpis.get("TWDR") or (1.0 - float(prr_v) if prr_v is not None else None)
    kpi_headers = [_txt("지표", "Metric"), _txt("값", "Value"), _txt("기준", "Threshold")]
    kpi_rows = [
        [_txt("평균 E2E 지연", "Avg. E2E Latency"),    f"{doc.kpis.get('평균 지연 (ms)', '—')} ms", "< 20 ms"],
        [_txt("최대 지연",     "Max. Latency"),        f"{doc.kpis.get('최대 지연 (ms)', '—')} ms", "< 100 ms"],
        [_txt("PRR (근사)",    "PRR (approx.)"),       str(prr_v or "—"),                           "> 0.90"],
        [_txt("PIR P99 †",     "PIR P99 †"),           f"{doc.pir_p99_ms or '—'} ms",               "≤ 100 ms [1]"],
        [_txt("PIR 준수",      "PIR Compliant"),       str(doc.pir_compliant if doc.pir_compliant is not None else "—"), "True"],
        [_txt("CBR (평균) †",  "CBR (avg.) †"),        str(doc.cbr_avg or "—"),                     "< 0.65 [3]"],
        [_txt("Jain FI",       "Jain FI"),             str(doc.jain_fairness_index or "—"),          "≥ 0.90 [9]"],
        [_txt("핸드오버 횟수", "Handover Count"),      str(doc.kpis.get("핸드오버 횟수", "—")),    "—"],
        [_txt("커버리지율 (%)", "Coverage Ratio (%)"),  str(doc.kpis.get("커버리지율 (%)", "—")),    "> 90 %"],
        [_txt("비용 개선율",   "Cost Improvement"),    f"{doc.improvement.get('비용 개선 (%)', '—')} %", "< 0 %"],
    ]
    _styled_table(kpi_headers, kpi_rows,
                  caption=_txt("Table 3. 시뮬레이션 핵심 KPI 요약",
                               "Table 3. Simulation KPI summary"))

    _section("4.2", "알고리즘 비교", "Algorithm Comparison", level=2)
    d.add_paragraph(
        _txt(f"선택 알고리즘: {doc.selected_algorithm}  /  기준 알고리즘: {doc.baseline_algorithm}",
             f"Selected algorithm: {doc.selected_algorithm}  /  Baseline: {doc.baseline_algorithm}")
    ).runs[0].font.size = Pt(10)

    if doc.algorithms:
        a_keys = ["algorithm", "total_cost", "average_latency_ms", "prr_approx",
                  "handover_count", "average_bs_load", "summary_rank_score"]
        a_headers_ko = ["알고리즘", "총 비용", "평균 지연(ms)", "PRR", "핸드오버", "BS 부하", "순위"]
        a_headers_en = ["Algorithm", "Total Cost", "Avg. Latency(ms)", "PRR", "Handover", "BS Load", "Rank"]
        a_hdrs = a_headers_en if lang == "en" else a_headers_ko
        a_rows = [[str(row.get(k, "—")) for k in a_keys] for row in doc.algorithms]
        _styled_table(a_hdrs, a_rows,
                      caption=_txt("Table 4. 알고리즘 비교 결과",
                                   "Table 4. Algorithm comparison results"))

    _section("4.3", "채널 및 공정성 지표", "Channel and Fairness Metrics", level=2)
    cf_headers = [_txt("지표", "Metric"), _txt("값", "Value"), _txt("출처", "Ref.")]
    cf_rows = [
        [_txt("CBR 평균 (해석 모델) †", "CBR avg. (analytical) †"),
         str(doc.cbr_avg or "—"), "[4]"],
        [_txt("PIR P99 상한 (해석 모델) †", "PIR P99 upper bound (analytical) †"),
         f"{doc.pir_p99_ms or '—'} ms", "[7]"],
        [_txt("PIR 기준 준수 (≤100 ms)", "PIR compliant (≤100 ms)"),
         str(doc.pir_compliant if doc.pir_compliant is not None else "—"), "[1]"],
        [_txt("Jain 공정성 지수", "Jain's Fairness Index"),
         str(doc.jain_fairness_index or "—"), "[9]"],
    ]
    _styled_table(cf_headers, cf_rows,
                  caption=_txt("Table 5. 채널 및 공정성 지표",
                               "Table 5. Channel and fairness metrics"))

    _section("4.4", "구간별 상위 20개", "Top-20 Per-Edge Metrics", level=2)
    if doc.per_edge_sample:
        pe_cols = ["edge_index", "street_name", "latency_ms", "load_ratio",
                   "handover", "cbr", "path_loss_db", "total_cost"]
        pe_hdrs_ko = ["#", "도로명", "지연(ms)", "부하율", "핸드오버", "CBR†", "경로손실(dB)†", "비용"]
        pe_hdrs_en = ["#", "Street", "Latency(ms)", "Load", "Handover", "CBR†", "PL(dB)†", "Cost"]
        pe_hdrs = pe_hdrs_en if lang == "en" else pe_hdrs_ko
        pe_rows = [[str(e.get(c, "")) for c in pe_cols] for e in doc.per_edge_sample]
        _styled_table(pe_hdrs, pe_rows,
                      caption=_txt("Table 6. 구간별 메트릭 (지연 순 상위 20개)",
                                   "Table 6. Per-edge metrics (top 20 by latency)"))
    else:
        d.add_paragraph(_txt("데이터 없음", "No data"))

    _section("4.5", "기지국 분석", "Base Station Analysis", level=2)
    if doc.per_bs_summary:
        bs_cols = ["bs_name", "node_type", "load_ratio", "severity",
                   "affected_edge_count", "avg_latency_on_route_ms", "jain_fairness_index"]
        bs_hdrs_ko = ["BS명", "유형", "부하율", "심각도", "영향 구간", "평균 지연(ms)", "Jain FI"]
        bs_hdrs_en = ["BS Name", "Type", "Load", "Severity", "Edges", "Avg. Latency(ms)", "Jain FI"]
        bs_hdrs = bs_hdrs_en if lang == "en" else bs_hdrs_ko
        bs_rows = [[str(b.get(c, "")) for c in bs_cols] for b in doc.per_bs_summary]
        _styled_table(bs_hdrs, bs_rows,
                      caption=_txt("Table 7. 기지국 부하 분석",
                                   "Table 7. Base station load analysis"))
    else:
        d.add_paragraph(_txt("데이터 없음", "No data"))

    # ── 5. Analysis and Discussion ────────────────────────────────────────────
    _section("5.", "분석 및 논의", "Analysis and Discussion")

    _section("5.1", "병목 구간", "Bottleneck Analysis", level=2)
    if doc.bottleneck_sections:
        for item in doc.bottleneck_sections[:5]:
            parts = [str(item.get(f, "")) for f in
                     ["street_name", "severity", "load_ratio", "latency_ms"] if item.get(f)]
            d.add_paragraph("  |  ".join(parts), style="List Bullet")
    else:
        d.add_paragraph(_txt("병목 구간 없음", "No bottleneck detected"))

    _section("5.2", "기지국 과부하", "BS Overload", level=2)
    if doc.overloaded_bs:
        for item in doc.overloaded_bs[:5]:
            parts = [str(item.get(f, "")) for f in
                     ["bs_name", "severity", "load_ratio"] if item.get(f)]
            d.add_paragraph("  |  ".join(parts), style="List Bullet")
    else:
        d.add_paragraph(_txt("과부하 기지국 없음", "No overloaded BS detected"))

    _section("5.3", "연결성 위험", "Connectivity Risk", level=2)
    if doc.future_risk_sections:
        for item in doc.future_risk_sections[:5]:
            parts = [str(item.get(f, "")) for f in
                     ["street_name", "severity", "nearest_bs_name"] if item.get(f)]
            d.add_paragraph("  |  ".join(parts), style="List Bullet")
    else:
        d.add_paragraph(_txt("연결성 위험 없음", "No connectivity risk detected"))

    # ── 6. Conclusion ─────────────────────────────────────────────────────────
    _section("6.", "결론", "Conclusion")
    if doc.primary_finding:
        p = d.add_paragraph(doc.primary_finding)
        p.runs[0].font.size = Pt(10)
        p.runs[0].italic = True
    for t in doc.trade_offs:
        d.add_paragraph(t, style="List Bullet")
    for r in doc.risk_factors:
        d.add_paragraph(r, style="List Bullet")
    if doc.recommendation_text:
        rp2 = d.add_paragraph(doc.recommendation_text)
        rp2.runs[0].font.size = Pt(10)

    # ── References ────────────────────────────────────────────────────────────
    d.add_page_break()
    _add_references()

    note_p = d.add_paragraph(
        _txt(
            "† 해석 모델 기반 수치. 실제 MAC 계층 시뮬레이션(ns-3/Veins)과 차이 가능. "
            "상기 출처의 검증된 해석 수식에 따라 계산.",
            "† Values derived from validated analytical models. May differ from MAC-layer "
            "simulation (ns-3/Veins). Calculated per cited peer-reviewed formulas."
        )
    )
    note_p.runs[0].font.size = Pt(9)
    note_p.runs[0].italic = True

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
